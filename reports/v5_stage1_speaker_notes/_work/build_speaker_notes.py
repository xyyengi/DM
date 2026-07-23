from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\DM_local")
OUT = ROOT / "reports" / "v5_stage1_speaker_notes" / "V5条件扩散模型汇报讲稿与甲方答疑手册.docx"
SOURCE = ROOT / "reports" / "v5_stage1_speaker_notes" / "_work" / "final_report_source.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GRAY = "666666"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "E2F0D9"
PALE_YELLOW = "FFF2CC"
PALE_RED = "FCE4D6"
RED = "9B1C1C"
INK = "222222"


def set_run_font(run, size=11, bold=False, italic=False, color=INK, east_asia="微软雅黑"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def fixed_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths_dxa, highlight=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    repeat_header(table.rows[0])
    prevent_split(table.rows[0])
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, 9.5, bold=True)
    for ri, values in enumerate(rows):
        row = table.add_row()
        prevent_split(row)
        for ci, value in enumerate(values):
            cell = row.cells[ci]
            if highlight and ri in highlight:
                shade_cell(cell, highlight[ri])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, 9.3, bold=(highlight and ri in highlight and ci == 0))
    fixed_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_script(doc, text, label="讲稿", color=INK):
    p = doc.add_paragraph(style="Speaker Script")
    r1 = p.add_run(f"【{label}】")
    set_run_font(r1, 11, bold=True, color=BLUE if label == "讲稿" else color)
    r2 = p.add_run(text)
    set_run_font(r2, 11, color=color)
    return p


def add_cue(doc, text):
    p = doc.add_paragraph(style="Speaker Cue")
    r = p.add_run(f"【屏幕动作】{text}")
    set_run_font(r, 9.5, italic=True, color=GRAY)
    return p


def add_note(doc, title, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    row = table.rows[0]
    repeat_header(row)
    prevent_split(row)
    cell = row.cells[0]
    shade_cell(cell, fill)
    set_cell_margins(cell, top=110, bottom=110, start=120, end=120)
    p = cell.paragraphs[0]
    r1 = p.add_run(title)
    set_run_font(r1, 10.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_run_font(r2, 10.3)
    fixed_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_qa(doc, question, short_answer, expanded):
    p = doc.add_paragraph(style="QA Question")
    r = p.add_run(question)
    set_run_font(r, 11.5, bold=True, color=DARK_BLUE)
    add_script(doc, short_answer, label="先答这一句", color=INK)
    if expanded:
        p2 = doc.add_paragraph(style="QA Detail")
        r1 = p2.add_run("展开：")
        set_run_font(r1, 10.5, bold=True, color=GRAY)
        r2 = p2.add_run(expanded)
        set_run_font(r2, 10.5, color=INK)


def add_page_number(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("第 ")
    set_run_font(r, 9, color=GRAY)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)
    r2 = p.add_run(" 页")
    set_run_font(r2, 9, color=GRAY)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    script = doc.styles.add_style("Speaker Script", 1)
    script.base_style = doc.styles["Normal"]
    script.paragraph_format.space_after = Pt(7)
    script.paragraph_format.line_spacing = 1.16

    cue = doc.styles.add_style("Speaker Cue", 1)
    cue.base_style = doc.styles["Normal"]
    cue.paragraph_format.left_indent = Inches(0.18)
    cue.paragraph_format.space_after = Pt(4)

    q = doc.styles.add_style("QA Question", 1)
    q.base_style = doc.styles["Normal"]
    q.paragraph_format.space_before = Pt(10)
    q.paragraph_format.space_after = Pt(3)
    q.paragraph_format.keep_with_next = True

    detail = doc.styles.add_style("QA Detail", 1)
    detail.base_style = doc.styles["Normal"]
    detail.paragraph_format.left_indent = Inches(0.18)
    detail.paragraph_format.space_after = Pt(7)
    detail.paragraph_format.line_spacing = 1.12


def build():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_page_number(section)
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("V5 条件扩散模型｜汇报讲稿与甲方答疑")
    set_run_font(hr, 9, color=GRAY)

    # Memo masthead
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("汇报讲稿与答疑手册")
    set_run_font(r, 23, bold=True, color="000000")
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r = p2.add_run("V5条件扩散模型 Stage-1 对比算例与消融实验")
    set_run_font(r, 14, color=GRAY)
    for label, value in [
        ("对应文档", "V5条件扩散模型Stage1对比算例与消融实验报告2.docx"),
        ("建议时长", "主讲约15分钟；问答按需展开"),
        ("核心结论", "当前推荐V5-TF Rank-1（epoch 8）；下一阶段先校准概率区间，再优化联合相关性"),
        ("使用方法", "“讲稿”可直接口述；“屏幕动作”和“答疑提示”不必念出"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rl = p.add_run(f"{label}：")
        set_run_font(rl, 10.5, bold=True)
        rv = p.add_run(value)
        set_run_font(rv, 10.5)

    add_note(
        doc,
        "汇报底线",
        "这次结果是Stage-1验证集结论，用于路线选择，不应表述为最终测试集泛化结论。不要说“V5-T证明时间步无效”，应说“只有时间步不足以替代forecast/calendar外部条件”。",
        PALE_YELLOW,
    )

    add_heading(doc, "一、60秒结论版", 1)
    add_script(
        doc,
        "本阶段在相同数据划分和采样协议下，对V4-RS、V5-T和V5-TF进行了验证集比较。结果显示，完整条件模型V5-TF综合表现最好：物理投影后，CRPS较V4-RS降低12.07%，ACF误差降低45.07%，净负荷MAE降低26.13%。V5-T只保留显式扩散时间步、没有forecast和calendar，结果未超过V4-RS，说明单靠时间步不足以替代外部条件。V5-TF通过独立条件编码和多尺度FiLM，把日前预测、日历和位置条件持续注入去噪网络，因此取得提升。当前还存在90%区间实际覆盖率只有81.60%、以及风光荷跨变量相关性未全面超过V4-RS的问题。下一阶段将先做不依赖测试集的概率区间校准，再针对联合相关结构进行诊断和模型优化，最后通过多随机种子复验后一次性评估测试集。",
    )

    add_heading(doc, "二、主讲稿（约15分钟）", 1)

    add_heading(doc, "1. 开场与阶段任务（约1分钟）", 2)
    add_cue(doc, "停留在“1. 结论先行”。")
    add_script(
        doc,
        "各位领导、专家好。今天汇报的是V5条件扩散模型的Stage-1对比算例。这个阶段的目标不是直接给出最终测试集结论，而是在严格统一的验证协议下回答三个问题：第一，旧版V4-RS的条件拼接方式是否还能改进；第二，显式扩散时间步和FiLM条件调制分别起什么作用；第三，生成场景能否同时满足概率质量、时间相关性和物理可行性。"
    )
    add_script(
        doc,
        "先说结论：当前推荐V5-TF Rank-1，也就是epoch 8。它在主要概率评分、时间相关性和净负荷指标上都优于V4-RS。物理越界问题已通过可审计的物理投影处理，不需要因为负光伏重新训练；但90%概率区间仍然欠覆盖，跨变量相关性也还没有全面超过基线，因此下一阶段工作很明确。"
    )

    add_heading(doc, "2. 数据、20条场景与采样协议（约2.5分钟）", 2)
    add_cue(doc, "切到“2. 实验内容与对比设置”和统一设置表。")
    add_script(
        doc,
        "三种模型采用相同的数据划分。训练集用于学习模型和拟合标准化参数；验证集用于选择checkpoint、生成场景和比较模型；测试集在这一阶段没有参与调参，留到模型方案完全锁定后再做最终评估。当前验证集有553个窗口，每个窗口包含未来168小时的风电、光伏和负荷。"
    )
    add_script(
        doc,
        "这里的20条场景，不是整个验证集只有20条，也不是20个批次。它的含义是：对每一个168小时窗口，都生成20种可能的风—光—荷联合轨迹。一条场景的形状是3乘168；一个窗口的场景集合是20乘3乘168。全部553个窗口合并后，共有553乘20，也就是11,060条多变量场景。"
    )
    add_script(
        doc,
        "当时终端里看到很多批次，是因为显存不能一次处理553个窗口。实际生成批大小是4，所以一次完整checkpoint评估需要向上取整553除以4，也就是139个生成批次。一个完整批次包含4个窗口，每个窗口20条场景，相当于同时处理80条轨迹；最后一个批次只剩1个窗口。批次只是计算和显存安排，不改变每个窗口20条场景的统计含义。"
    )
    add_script(
        doc,
        "固定生成随机种子424242，是为了让伪随机噪声序列可复现，并尽量让三种模型面对同一套采样随机条件。相同种子不会让三种模型输出相同，因为模型参数和映射不同；它只是减少抽样偶然性。这里还要区分：训练种子是2026，生成种子是424242。当前只完成了一个训练种子的Stage-1比较，下一阶段仍需多训练种子复验。"
    )
    add_script(
        doc,
        "500步反向采样指每条场景从高斯噪声开始，依次从x500去噪到x0，共执行500个反向扩散步骤。它不是预测500小时，也不是训练500轮。最终得到的是残差轨迹，再用“日前预测减生成残差”还原风电、光伏和负荷场景。"
    )
    add_note(
        doc,
        "如果甲方追问为什么只用20条",
        "回答：20条是Stage-1在概率表达能力与计算量之间的统一折中，足以完成三模型初步比较，但尾部分位数仍较粗。下一阶段会在冻结模型后补做20、50、100条场景敏感性分析，不一定需要重训。",
        PALE_BLUE,
    )

    add_heading(doc, "3. 三种模型和条件注入区别（约3分钟）", 2)
    add_cue(doc, "切到“三组对比实验”表。")
    add_script(
        doc,
        "V4-RS是旧版工程基线。它把3通道带噪状态、3通道日前预测和8通道时间编码在输入端直接拼接，形成14通道输入，再统一送入U-Net。也就是说，状态和条件从第一层开始就混在一起，由网络自己判断哪些是待去噪状态、哪些是辅助条件。"
    )
    add_script(
        doc,
        "V5-T只输入3通道带噪状态，同时显式加入真实扩散时间步。它的作用是检查：只让网络知道当前处于第几个去噪步骤，是否足以改进结果。要注意，V5-T不是在V4-RS上只增加一个时间步；它同时去掉了forecast和calendar。因此V5-T低于V4-RS不能解释成时间步无效，只能说明只有时间步不足以替代外部条件信息。"
    )
    add_script(
        doc,
        "V5-TF把状态主干和条件支路分开。状态主干始终只输入3通道带噪状态；日前预测、8通道日历特征和相对位置分别经过独立编码，然后只在条件支路内部融合，并下采样成多个时间尺度的条件特征。这些条件特征不作为额外通道拼到状态入口，而是在编码器、瓶颈和解码器的各级残差块中生成FiLM参数。"
    )
    add_script(
        doc,
        "FiLM可以写成h撇等于1加gamma，乘归一化后的h，再加beta。gamma决定每个特征通道放大或抑制多少，beta决定向上或向下平移多少。扩散时间步产生一组gamma和beta，forecast、calendar、位置条件再产生一组，最后相加。直观上，拼接是把条件当作输入数据一起卷积；FiLM注入是让条件成为控制信号，在网络多个尺度持续告诉主干哪些特征应该加强、减弱或偏移。"
    )
    add_note(
        doc,
        "代码层面的准确说法",
        "V5-TF并非代码里完全没有concat：forecast、calendar和位置在各自编码后，会在“条件编码器内部”融合；U-Net解码器也会拼接skip connection。真正的区别是：条件没有与x_t在状态主干入口拼成14通道。",
        PALE_YELLOW,
    )

    add_heading(doc, "4. 综合指标结果（约2分钟）", 2)
    add_cue(doc, "切到“3. 主要结果”、结果表和图1。")
    add_script(
        doc,
        "物理投影后的Rank-1结果显示，V5-TF的CRPS是728.39，较V4-RS降低12.07%；MVES降低11.29%；ACF误差降低45.07%；净负荷MAE降低26.13%；6小时爬坡MAE降低4.32%。这些指标均为越低越好，因此V5-TF不是只在一个指标上领先，而是在概率质量、时间结构和净负荷应用指标上形成了比较一致的改善。"
    )
    add_script(
        doc,
        "V5-T的90%覆盖率是82.67%，看起来略高于V5-TF的81.60%，但它的90%区间宽度也最大，CRPS、MVES和爬坡误差都更差。这说明覆盖率不能脱离区间宽度单独看；区间拉得很宽也能提高覆盖，但未必代表概率分布质量更好。"
    )

    add_heading(doc, "5. 典型窗口和场景包络（约1.5分钟）", 2)
    add_cue(doc, "切到“4. 场景生成效果”、图2和图3。")
    add_script(
        doc,
        "这里的典型验证窗口不是人工看完曲线后挑选，也不是选择V5-TF表现最好的窗口。具体做法是：先对553个验证窗口分别计算168小时日前净负荷预测MAE。净负荷定义为负荷减风电减光伏。然后计算553个MAE的中位数，选择与该中位数最接近的窗口。最终选中第436个窗口，程序零基索引是435，该窗口净负荷日前预测MAE为2714.49兆瓦。"
    )
    add_script(
        doc,
        "这个规则只使用共同的日前预测和实测值，不使用三种模型的生成结果，所以不会偏向任何模型。这里的“典型”准确地说是预测难度接近验证集中位水平，不代表它在所有气象或极端事件特征上都最典型。图2比较同一窗口下三种模型的90%包络；图3进一步展示V5-TF的20条场景、场景中位数和包络。"
    )

    add_heading(doc, "6. 时间特性、消融与物理投影（约2分钟）", 2)
    add_cue(doc, "依次切到第5节图4、第6节两个表。")
    add_script(
        doc,
        "时间相关性方面，V5-TF整体ACF误差最低，光伏和负荷改善最明显，6小时净负荷爬坡也优于V4-RS。但跨变量相关误差方面，V5-TF是0.06487，仍高于V4-RS的0.04346。这说明V5-TF已经较好地学到了每个变量自身的时间结构，但风、光、荷三者之间的联合依赖还没有全面超过旧基线。"
    )
    add_script(
        doc,
        "条件消融进一步说明了外部条件的重要性。去掉日历后CRPS上升11.06%；去掉日前预测后上升38.54%，而且时间相关性和跨变量相关性同时恶化。因此日前预测是最关键的外部条件，日历也提供了稳定的周期信息。"
    )
    add_script(
        doc,
        "原始输出中的负风电和负光伏并不是数据清洗遗漏。训练样本虽然已经非负，但扩散模型在无界的标准化残差空间随机生成，经过“实际值等于预测值减残差”重构后，夜间forecast接近零时仍可能产生负光伏。当前物理投影是在反标准化和场景重构后，把风电、光伏截断到零和容量上限之间，把负荷限制为非负，并在光伏预测小于等于1兆瓦时置零。它不修改模型权重，原始数组也保留用于审计。"
    )

    add_heading(doc, "7. 下一阶段怎么做（约3分钟）", 2)
    add_cue(doc, "切到“7. 阶段结论与下一步”。")
    add_script(
        doc,
        "下一阶段第一项是概率区间校准。当前V5-TF的名义90%区间实际覆盖率是81.60%，属于欠覆盖。这里优先采用后处理校准，不必立刻重新训练模型。首先冻结当前V5-TF权重和物理投影规则；其次补做20、50、100条场景敏感性分析，确认尾部分位数是否随场景数稳定；然后划出独立校准数据，不能使用最终测试集，根据真实值落在上下分位数之外的程度学习一个有限样本校准量。"
    )
    add_script(
        doc,
        "具体可以采用conformal，也就是保形校准。对名义90%区间，先得到5%和95%分位数；再在校准集上计算每个真实值超出区间的距离，取符合90%有限样本要求的分位数q星，把区间统一或分组扩展为下界减q星、上界加q星。风电、光伏、负荷量纲和误差水平不同，建议分别校准，并进一步按预测时效、光伏昼夜和高爬坡时段检查条件覆盖率。目标不是盲目拉宽区间，而是在覆盖率接近80%、90%、95%的同时，控制区间宽度和CRPS。"
    )
    add_script(
        doc,
        "第二项是风光荷联合相关性优化。先做诊断，不马上堆结构：分别检查风—光、风—荷、光—荷三对相关误差，按预测时效、日夜、季节和高爬坡窗口分层，并比较物理投影前后相关矩阵，找出误差来源。短期可以尝试基于历史残差依赖的经验copula或ensemble reordering，在不改变各变量边际分布的前提下重排场景成员，验证是否能先修复联合依赖。"
    )
    add_script(
        doc,
        "如果后处理不足，再进行结构性重训：在V5-TF瓶颈加入显式跨变量注意力或关系建模；在日前预测条件中增加总新能源、净负荷和爬坡特征；训练目标从单一噪声MSE扩展为噪声损失加小权重的相关矩阵损失、净负荷损失和爬坡损失。相关损失可以比较每个168小时窗口内预测洁净残差与真实残差的三乘三相关矩阵，避免只优化单变量。所有新增项都要逐项消融，防止改善相关性却损害CRPS或覆盖率。"
    )
    add_script(
        doc,
        "最后，在概率校准和联合相关性方案确定后，至少使用3个训练随机种子复验，预先固定模型选择规则和评价指标，再一次性运行测试集。这样最终结论才能同时回答平均性能、随机种子稳定性和独立测试泛化。"
    )

    add_heading(doc, "8. 收尾（约30秒）", 2)
    add_script(
        doc,
        "总结来说，V5-TF已经证明独立条件编码和多尺度FiLM能够显著改善场景概率质量和时间结构，物理越界也已经通过可审计投影解决。当前不需要因为负值重新训练，但还不能把验证结果当作最终测试结论。下一阶段将围绕“区间校准、联合相关性、多随机种子和锁定测试集”四件事推进。我的汇报结束，请各位专家指正。"
    )

    doc.add_page_break()
    add_heading(doc, "三、甲方高频问答", 1)
    add_note(
        doc,
        "答疑原则",
        "先用一到两句话正面回答，再按对方兴趣展开。遇到尚未做的工作要说“当前验证结果尚不能单独证明”，不要用推测替代实验结论。",
        PALE_GREEN,
    )

    add_qa(
        doc,
        "1. “20条场景”到底是什么意思？",
        "是每一个168小时验证窗口生成20条可能的风—光—荷联合轨迹，不是整个验证集只有20条。",
        "验证集有553个窗口，因此总计553×20=11,060条场景；内部数组为[553,20,3,168]，展平后为[11060,3,168]。20条共同近似条件概率分布，用于计算中位数、包络、覆盖率、CRPS和Energy Score。",
    )
    add_qa(
        doc,
        "2. 为什么终端里会有139个批次？",
        "139是显存受限下的数据生成批次数，和每个窗口20条场景不是同一个概念。",
        "生成batch size为4，553个验证窗口需要ceil(553/4)=139批。每个完整批次处理4个窗口×20条场景=80条轨迹；最后一批处理1个窗口。训练batch size为64、生成batch size为4、每窗场景数为20，三者含义不同。",
    )
    add_qa(
        doc,
        "3. 这些场景来自验证集还是测试集？",
        "本报告全部Stage-1场景和指标来自验证集，测试集没有参与模型选择。",
        "训练集用于模型训练和标准化；验证集用于选checkpoint、生成和比较；测试集在方案锁定后一次性使用。这样避免根据测试结果反复改模型造成信息泄漏。",
    )
    add_qa(
        doc,
        "4. 固定随机种子有什么作用？",
        "固定种子使随机采样可复现，并减少三种模型比较中的抽样偶然性。",
        "生成种子是424242；训练种子是2026。相同种子不会让不同模型输出相同，因为网络参数和映射不同。固定一个生成种子也不等于完成了模型稳定性验证，因此下一阶段仍需多个训练种子。",
    )
    add_qa(
        doc,
        "5. 500步反向采样是什么意思？",
        "每条场景从高斯噪声开始，经过500次逐级去噪形成残差轨迹；500不是预测小时数，也不是训练轮数。",
        "过程可理解为x500→x499→…→x0。每一步网络都根据当前带噪状态、扩散步和条件预测去噪方向。20条场景通常批量并行，但计算量仍显著高于一次普通前向预测。",
    )
    add_qa(
        doc,
        "6. 你们说的“基线”是什么？",
        "基线就是旧版工程模型V4-RS，用来判断完整V5方案是否真正优于现有方案。",
        "V4-RS采用x_t、forecast和calendar入口拼接；V5-T采用显式时间步但无forecast/calendar；V5-TF采用显式时间步和多尺度FiLM条件。V4-RS与V5-T不是严格单变量消融，因此不能据此说时间步无效。",
    )
    add_qa(
        doc,
        "7. 为什么V5-T没有超过V4-RS？",
        "因为V5-T虽然补上了显式扩散时间步，但同时去掉了forecast和calendar；时间步收益不足以弥补外部条件信息的损失。",
        "准确结论是“只有时间步不够”，而不是“时间步没有用”。若要单独量化时间步净贡献，还应增加保留forecast/calendar与FiLM、仅关闭t分支的桥接消融。",
    )
    add_qa(
        doc,
        "8. 拼接和FiLM条件注入有什么区别？",
        "拼接是把状态和条件作为14个输入通道一起卷积；FiLM是状态走3通道主干，条件单独编码后在多层生成缩放和偏移参数控制主干特征。",
        "V5-TF中forecast、calendar和位置先分别编码，只在条件支路内部融合；随后通过1×1卷积产生gamma和beta，在各残差块执行h'=(1+gamma)×Norm(h)+beta。因此条件是控制信号，而不是主干入口的待去噪数据。",
    )
    add_qa(
        doc,
        "9. 典型验证窗口是怎么选的？",
        "先计算553个窗口的168小时日前净负荷预测MAE，再选择与全体中位数最接近的窗口；模型生成结果不参与选择。",
        "净负荷=负荷−风电−光伏。选中第436个窗口，零基索引435，MAE为2714.49 MW。它代表中等预测难度，不是最好看、最简单或V5-TF最优的窗口，也不代表所有气象特征都最典型。",
    )
    add_qa(
        doc,
        "10. 数据清洗后为什么还会生成负光伏？",
        "清洗约束的是训练样本，不能自动约束扩散模型的无界随机输出；夜间forecast接近零时，残差重构仍可能得到负值。",
        "模型生成标准化残差，实际场景=forecast−residual。扩散采样从高斯噪声出发，输出空间是实数域，因此非负训练数据不等于生成结果天然非负。",
    )
    add_qa(
        doc,
        "11. 物理投影是什么？需要重新训练吗？",
        "物理投影是生成后的可行域映射，不修改模型权重，因此不需要为此重新训练。",
        "风电和光伏截断到[0,容量上限]，负荷限制为非负，光伏forecast≤1 MW时置零。原始结果不覆盖，另存投影结果用于工程分析，保证业务可用性和科研审计同时成立。",
    )
    add_qa(
        doc,
        "12. 90%区间覆盖率为什么只有81.60%，下一步怎么校准？",
        "说明当前90%包络偏窄或尾部离散度不足；下一步优先做独立校准集上的保形校准，不使用测试集调参。",
        "冻结V5-TF后，计算5%和95%分位区间；在独立校准集上用真实值超出区间的距离确定q*，得到[下界−q*,上界+q*]。风、光、荷分开校准，并按预测时效、昼夜和高爬坡时段检查条件覆盖。校准同时比较覆盖率、宽度和CRPS，避免单纯拉宽。",
    )
    add_qa(
        doc,
        "13. 风光荷联合相关性下一步具体怎么优化？",
        "先定位哪一对变量、哪个时效和工况出问题；短期尝试经验copula重排，若不足再用跨变量注意力和相关性辅助损失重训。",
        "诊断风—光、风—荷、光—荷相关矩阵及投影前后差异；结构上可增加跨变量注意力、净负荷和总新能源条件；损失上加入小权重相关矩阵、净负荷及爬坡损失。所有改动需消融验证，保证改善相关性时不牺牲CRPS和覆盖率。",
    )
    add_qa(
        doc,
        "14. 为什么不现在就报告测试集？",
        "因为模型和校准方案还在选择阶段，现在查看测试集会把测试信息带回开发过程，削弱最终独立性。",
        "正确流程是验证集完成架构、损失、采样规模和校准方案选择；至少3个训练随机种子复验；预先固定评价规则；最后一次性运行测试集。",
    )
    add_qa(
        doc,
        "15. 目前结果能否直接用于工程？",
        "投影后的场景已具备物理可行性，可用于阶段性分析；但在正式风险决策前仍需完成概率校准、联合相关性优化和独立测试。",
        "当前适合做项目阶段汇报、方案比较和业务流程联调，不宜把81.60%的90%覆盖区间解释为已经完全校准的90%置信范围。",
    )

    doc.add_page_break()
    add_heading(doc, "四、下一阶段技术路线速查", 1)
    add_table(
        doc,
        ["阶段", "核心动作", "是否重训", "验收指标"],
        [
            ["A. 采样稳定性", "冻结V5-TF；比较20/50/100成员；检查分位数和CRPS稳定性", "否", "指标随成员数趋稳；选择计算成本可接受的最小规模"],
            ["B. 概率校准", "独立校准集；分变量保形校准；按时效/昼夜/爬坡分层诊断", "通常否", "80/90/95%覆盖接近名义值，宽度和CRPS可控"],
            ["C. 联合依赖诊断", "分解三对相关误差；检查投影前后、时效和事件分层", "否", "明确误差来源和优先变量对"],
            ["D. 依赖后处理", "经验copula或ensemble reordering，保持边际分布", "否", "跨变量相关误差下降且CRPS不恶化"],
            ["E. 结构性优化", "跨变量注意力；净负荷/新能源条件；相关、爬坡辅助损失", "是", "相关误差优于V4-RS，同时保持V5-TF概率优势"],
            ["F. 最终确认", "至少3个训练种子；固定规则；一次性测试集评估", "是", "均值、方差、配对置信区间和测试泛化均通过"],
        ],
        [1250, 4300, 1200, 2610],
        highlight={1: PALE_BLUE, 4: PALE_YELLOW, 5: PALE_GREEN},
    )

    add_heading(doc, "五、汇报中不建议使用的表述", 1)
    add_table(
        doc,
        ["不建议说", "建议改为"],
        [
            ["“V5-T证明扩散时间步没有用。”", "“V5-T说明只有时间步不足以替代forecast/calendar；当前对比不能单独量化时间步净贡献。”"],
            ["“90%区间已经达到90%可信度。”", "“当前名义90%区间实际覆盖81.60%，下一阶段需要独立校准。”"],
            ["“典型窗口代表整个数据集所有特征。”", "“该窗口代表日前净负荷预测难度接近验证集中位水平。”"],
            ["“物理投影修复了训练数据问题。”", "“物理投影约束生成输出；负值不是清洗遗漏。”"],
            ["“V5-TF最终优于所有模型。”", "“V5-TF在当前Stage-1验证协议下综合最优，最终结论待多种子和测试集确认。”"],
        ],
        [3600, 5760],
    )

    doc.core_properties.title = "V5条件扩散模型汇报讲稿与甲方答疑手册"
    doc.core_properties.subject = "Stage-1对比算例主讲稿、常见问答与下一阶段技术路线"
    doc.core_properties.author = "DM项目组"
    doc.core_properties.keywords = "V5-TF, FiLM, 场景生成, 概率校准, 联合相关性, 汇报讲稿"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)
    print("size", OUT.stat().st_size)
    print("source", SOURCE)


if __name__ == "__main__":
    build()
