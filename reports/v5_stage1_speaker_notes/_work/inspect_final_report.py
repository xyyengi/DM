from pathlib import Path
import json
import zipfile
from docx import Document

source = Path(r"D:\DM_local\reports\v5_stage1_speaker_notes\_work\final_report_source.docx")
out = source.with_name("final_report_content.json")
doc = Document(source)

paragraphs = []
for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    if text:
        paragraphs.append({
            "index": i,
            "style": paragraph.style.name,
            "text": text,
        })

tables = []
for ti, table in enumerate(doc.tables):
    rows = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])
    tables.append({"index": ti, "rows": rows})

with zipfile.ZipFile(source) as archive:
    parts = archive.namelist()
    media = [name for name in parts if name.startswith("word/media/")]

record = {
    "paragraph_count": len(doc.paragraphs),
    "nonempty_paragraph_count": len(paragraphs),
    "table_count": len(doc.tables),
    "image_count": len(doc.inline_shapes),
    "section_count": len(doc.sections),
    "paragraphs": paragraphs,
    "tables": tables,
    "media": media,
}
out.write_text(json.dumps(record, ensure_ascii=False, indent=2), encoding="utf-8")
print(json.dumps({k: record[k] for k in record if k not in {"paragraphs", "tables"}}, ensure_ascii=False, indent=2))
