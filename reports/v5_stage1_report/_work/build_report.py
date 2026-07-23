from __future__ import annotations

import csv
import json
import shutil
from pathlib import Path

import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\DM_local")
WORK = ROOT / "reports" / "v5_stage1_report" / "_work"
FIG = WORK / "figures"
REFERENCE = WORK / "reference.docx"
OUTPUT = ROOT / "reports" / "v5_stage1_report" / "V5条件扩散模型Stage1对比算例与消融实验报告.docx"
AUDIT_JSON = WORK / "report_evidence.json"
CSV_PATH = (
    ROOT
    / "outputs_shandong"
    / "v5_stage1"
    / "comparisons"
    / "20260723_physical_projection"
    / "v5_stage1_comparison.csv"
)

RESULTS = {
    "V4-RS": ROOT
    / "outputs_shandong"
    / "v5_stage1"
    / "20260722_143437_v4rs_repro_stage1_seed2026_20260722_143431_val_rank1_epoch11_posterior_n20_seed424242",
    "V5-T": ROOT
    / "outputs_shandong"
    / "v5_stage1"
    / "20260722_151755_v5_t_stage1_seed2026_20260722_151749_val_rank1_epoch29_posterior_n20_seed424242",
    "V5-TF": ROOT
    / "outputs_shandong"
    / "v5_stage1"
    / "20260722_155013_v5_tf_stage1_seed2026_20260722_155007_val_rank1_epoch8_posterior_n20_seed424242",
}

ARCH = {"v4_legacy": "V4-RS", "v5_t": "V5-T", "v5_tf": "V5-TF"}
BLUE = "D9EAF7"
LIGHT_BLUE = "EAF3F8"
LIGHT_GRAY = "F2F2F2"
PALE_GREEN = "E2F0D9"
PALE_RED = "FCE4D6"


def read_csv() -> list[dict[str, str]]:
    with CSV_PATH.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def rank1_rows(rows: list[dict[str, str]]) -> dict[str, dict[str, str]]:
    out = {}
    for row in rows:
        if row["checkpoint_rank"] == "1" and row["condition_ablation"] == "none":
            out[ARCH[row["architecture"]]] = row
    return out


def constrained_metrics() -> dict[str, dict]:
    return {
        name: json.loads((path / "metrics_constrained.json").read_text(encoding="utf-8"))
        for name, path in RESULTS.items()
    }


def per_window_crps(actual: np.ndarray, scenarios: np.ndarray) -> np.ndarray:
    # Ensemble CRPS = mean|X-y| - 0.5 E|X-X'|, evaluated pointwise then
    # averaged across the three variables and 168 h in each validation window.
    n_members = scenarios.shape[1]
    term1 = np.mean(np.abs(scenarios - actual[:, None, :, :]), axis=1)
    xs = np.sort(scenarios, axis=1)
    weights = (2 * np.arange(1, n_members + 1) - n_members - 1).reshape(1, -1, 1, 1)
    term2 = np.sum(xs * weights, axis=1) / (n_members**2)
    return np.mean(term1 - term2, axis=(1, 2))


def bootstrap_evidence() -> dict:
    actual = np.load(RESULTS["V5-TF"] / "actual_data.npy")
    per_window = {}
    for name, path in RESULTS.items():
        scenarios = np.load(path / "actual_scenarios_constrained.npy", mmap_mode="r")
        per_window[name] = per_window_crps(actual, scenarios)

    rng = np.random.default_rng(20260723)
    indices = rng.integers(0, actual.shape[0], size=(10000, actual.shape[0]))
    out = {}
    for comparator in ["V4-RS", "V5-T"]:
        delta = per_window["V5-TF"] - per_window[comparator]
        boot = delta[indices].mean(axis=1)
        out[f"V5-TF_vs_{comparator}"] = {
            "mean_delta": float(delta.mean()),
            "ci95": [float(x) for x in np.quantile(boot, [0.025, 0.975])],
            "window_better_pct": float(np.mean(delta < 0) * 100),
            "n_windows": int(delta.size),
            "bootstrap_repetitions": 10000,
            "seed": 20260723,
        }
    return out


def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_font(run, size=12, bold=False, color=None, east_asia="宋体") -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=70, bottom=70, end=70) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None, highlight_rows=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, 9, bold=True)
    highlight_rows = highlight_rows or {}
    for ridx, row_data in enumerate(rows):
        new_row = table.add_row()
        prevent_row_split(new_row)
        cells = new_row.cells
        fill = highlight_rows.get(ridx)
        for cidx, text in enumerate(row_data):
            cell = cells[cidx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if fill:
                set_cell_shading(cell, fill)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if cidx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(text))
            set_run_font(r, 8.5, bold=(ridx in highlight_rows and cidx == 0))
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_body(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.33)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, 11.5, bold=True)
        r2 = p.add_run(text[len(bold_lead) :])
        set_run_font(r2, 11.5)
    else:
        r = p.add_run(text)
        set_run_font(r, 11.5)


def add_equation(doc: Document, equation: str, number: int, note: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_together = True
    r = p.add_run(f"{equation}    （{number}）")
    r.font.name = "Cambria Math"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    r.font.size = Pt(11.5)
    if note:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(5)
        r2 = p2.add_run(note)
        set_run_font(r2, 9.5, color="666666")


def add_callout(doc: Document, title: str, text: str, fill=PALE_GREEN) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, 11, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(text)
    set_run_font(r2, 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(level=level)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, 17 if level == 1 else 13, bold=True, east_asia="微软雅黑")


def add_picture(doc: Document, filename: str, caption: str, width=6.35) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(FIG / filename), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption.split("　", 1)[0])
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(5)
    r = cap.add_run(caption)
    set_run_font(r, 9.5, color="555555")


def add_page_number(section) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("第 ")
    set_run_font(r, 9, color="777777")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)
    r2 = p.add_run(" 页")
    set_run_font(r2, 9, color="777777")


def pct_reduction(base: float, new: float) -> float:
    return (base - new) / base * 100


def fmt(v, digits=2) -> str:
    return f"{float(v):,.{digits}f}"


def main() -> None:
    rows = read_csv()
    rank1 = rank1_rows(rows)
    cm = constrained_metrics()
    bootstrap = bootstrap_evidence()
    fig_summary = json.loads((FIG / "figure_summary.json").read_text(encoding="utf-8"))

    evidence = {
        "comparison_csv": str(CSV_PATH),
        "rank1": rank1,
        "constrained_metrics": cm,
        "paired_bootstrap": bootstrap,
        "figure_summary": {k: v for k, v in fig_summary.items() if k != "acf_curves"},
    }
    AUDIT_JSON.write_text(json.dumps(evidence, ensure_ascii=False, indent=2), encoding="utf-8")

    shutil.copy2(REFERENCE, OUTPUT)
    doc = Document(OUTPUT)
    clear_body(doc)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(11.5)
    for section in doc.sections:
        section.top_margin = Inches(0.78)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.82)
        section.right_margin = Inches(0.82)
        add_page_number(section)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(68)
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("V5 条件扩散模型")
    set_run_font(r, 25, bold=True, east_asia="微软雅黑")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("Stage-1 对比算例与消融实验报告")
    set_run_font(r, 21, bold=True, east_asia="微软雅黑")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(12)
    r = p3.add_run("V4-RS / V5-T / V5-TF · 山东数据集 · 验证集阶段")
    set_run_font(r, 12.5, color="555555", east_asia="微软雅黑")
    doc.add_paragraph().paragraph_format.space_after = Pt(50)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["报告性质", "模型开发阶段的对比算例、消融实验与物理可行性审计"],
            ["数据范围", "验证集 553 个窗口；每窗 168 h；风电/光伏/负荷 3 变量"],
            ["场景设置", "20 条场景；posterior 反向方差；500 步采样；固定 seed=424242"],
            ["结果口径", "主表采用物理投影后结果；原始生成结果单独保留并列示"],
            ["版本", "代码提交 7321265d；报告日期 2026-07-23"],
        ],
        widths=[1.3, 4.8],
    )
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(50)
    r = p4.add_run("项目汇报用")
    set_run_font(r, 12, color="777777", east_asia="微软雅黑")
    p4.add_run().add_break(WD_BREAK.PAGE)

    add_heading(doc, "摘要与结论先行", 1)
    add_callout(
        doc,
        "结论：当前不需要因“太阳能小于 0”重新训练。",
        "负值来自扩散模型在无界标准化残差空间中的随机生成，而不是数据清洗遗漏。现已增加可审计的物理可行域投影：保留原始输出，同时生成约束后场景。约束后所有模型的物理越界率均为 0%，因此可以进入 Stage-1 结果分析与项目汇报。",
    )
    add_body(
        doc,
        "在统一训练/验证划分、统一 20 场景、统一采样种子与 posterior 采样协议下，V5-TF Rank-1（epoch 8）取得最佳综合结果。物理投影后，其总 CRPS 为 728.39，较 V4-RS 降低 "
        f"{pct_reduction(float(rank1['V4-RS']['constrained_total_crps']), float(rank1['V5-TF']['constrained_total_crps'])):.2f}%；"
        "MVES 为 23,970.88；ACF 误差为 0.01096；净负荷 MAE 为 1,943.59 MW。"
    )
    add_body(
        doc,
        "V5-T 只修正扩散时间步条件，但不使用 forecast/calendar 外部条件，结果并未优于 V4-RS。V5-TF 在此基础上加入多尺度 forecast/calendar 编码与 FiLM 调制后明显领先，说明性能增益主要来自“正确扩散时间步 + 外部条件调制”的组合，而不是简单增加模型参数。"
    )
    add_body(
        doc,
        "当前结论属于 Stage-1 验证集结论，足以用于模型路线选择和项目阶段汇报，但不应写成最终测试集泛化结论。正式定版前建议补做多随机种子复验，再锁定模型并一次性评估测试集。"
    )

    add_heading(doc, "1. 报告定位：模型、实验与算例的区别", 1)
    add_body(
        doc,
        "模型是可训练的数学结构及其推理规则；实验是为了验证某个假设而设计的受控流程；算例是把模型和实验落到一组具体数据、参数、工况与结果上的完整实例。本报告可称为“V5 Stage-1 对比算例”，其中包含一组主对比实验、两组条件消融实验，以及一组物理投影前后审计。"
    )
    add_table(
        doc,
        ["概念", "本项目中的对应内容", "作用"],
        [
            ["模型", "V4-RS、V5-T、V5-TF 三种条件扩散架构", "规定输入、条件注入方式和去噪网络"],
            ["实验", "统一协议下的训练、Top-3 验证生成、条件消融、投影前后比较", "隔离变量并验证模型设计是否有效"],
            ["算例", "山东验证集 553×168 h、20 场景、固定种子及完整指标/曲线", "形成可复现、可展示、可审计的项目结果"],
        ],
        widths=[0.7, 3.2, 2.2],
    )

    add_heading(doc, "2. 数据、任务与统一评估协议", 1)
    add_body(
        doc,
        "三种模型均学习日前预测误差的条件分布。残差定义为 residual = forecast − actual；扩散模型生成残差场景后，按 actual_scenario = forecast − residual_scenario 重构风电、光伏和负荷场景。这样可以把确定性日前预测作为中心信息，让生成模型重点学习预测误差的不确定性。"
    )
    add_equation(
        doc,
        "r₀ = f − y，    ŷ⁽ˢ⁾ = f − r̂₀⁽ˢ⁾",
        1,
        "其中 f 为日前预测，y 为实测值，r₀ 为真实残差，s 表示第 s 条随机场景。",
    )
    add_table(
        doc,
        ["项目", "统一设置"],
        [
            ["数据切分", "仅使用验证集进行 Stage-1 比较；测试集未参与模型选择"],
            ["验证规模", "553 个独立窗口，每个窗口 168 h，3 个变量"],
            ["生成规模", "每个窗口 20 条场景，共 11,060 条多变量场景"],
            ["采样", "posterior 方差，500 个反向扩散步，generation seed=424242"],
            ["Checkpoint", "每个架构按 validation epsilon-MSE 选 Top-3，再统一生成评估"],
            ["主报告口径", "Rank-1 + 物理投影后；原始输出和 Rank-2/3 留作稳定性审计"],
        ],
        widths=[1.35, 4.8],
    )

    add_heading(doc, "3. 三种架构及新增原理", 1)
    add_heading(doc, "3.1 V4-RS：旧版残差扩散基线", 2)
    add_body(
        doc,
        "V4-RS 使用 14 通道拼接输入：当前噪声状态 3 通道、日前预测 3 通道、日历/时间编码 8 通道。条件信息直接拼接到 U-Net 输入。实现审计表明，该旧版去噪器没有把扩散时间步 t 作为独立输入显式送入残差块；因此日历时间特征存在，但“当前处于第几个扩散去噪步骤”的信息表达不充分。"
    )
    add_equation(
        doc,
        "u_V4 = Concat(xₜ, f, k) ∈ ℝ¹⁴×ᴸ",
        2,
        "xₜ、f、k 分别为 3 通道噪声状态、3 通道日前预测和 8 通道日历编码；L=168。",
    )
    add_heading(doc, "3.2 V5-T：只补齐扩散时间步条件", 2)
    add_body(
        doc,
        "V5-T 的状态输入保持 3 通道，仅接收带噪目标 x_t；扩散时间步 t 经过正弦位置编码和 MLP，再调制各残差块。该模型不接收 forecast 和 calendar，用来隔离检验“正确扩散时间步条件”本身的作用。结果表明，只加入 t 并不足以获得最优场景。"
    )
    add_equation(
        doc,
        "eₜ,₂ₖ = sin(t / 10000²ᵏ⁄ᵈ)，    eₜ,₂ₖ₊₁ = cos(t / 10000²ᵏ⁄ᵈ)",
        3,
        "正弦/余弦编码把离散扩散步 t 映射为连续向量，再经 MLP 生成各残差块所需的调制参数。",
    )
    add_heading(doc, "3.3 V5-TF：时间步 + forecast/calendar 的 FiLM 条件调制", 2)
    add_body(
        doc,
        "V5-TF 将 forecast 3 通道、calendar 8 通道和相对位置分别编码，并在编码器、瓶颈和解码器的多个尺度融合。条件不再粗暴拼接进状态主干，而是通过 FiLM（Feature-wise Linear Modulation，逐特征线性调制）控制中间特征："
    )
    add_equation(
        doc,
        "c = Fuse(E_f(f), E_cal(k), E_pos(p))",
        4,
        "E_f、E_cal、E_pos 分别编码日前预测、日历特征和 168 h 相对位置，Fuse 表示多尺度融合。",
    )
    add_equation(
        doc,
        "γ = γₜ + γ_c，    β = βₜ + β_c",
        5,
        "时间步分支与外部条件分支分别产生缩放和偏移参数，再做加性融合。",
    )
    add_equation(
        doc,
        "h′ = (1 + γ) ⊙ Norm(h) + β",
        6,
        "γ 控制逐通道缩放，β 控制逐通道平移，⊙ 为逐元素乘法；该操作部署在编码器、瓶颈和解码器。",
    )
    add_body(
        doc,
        "这一设计把“扩散过程走到哪一步”和“当前气象/日历条件是什么”分开编码，再在多尺度特征上组合，既保持 3 通道状态空间语义稳定，也让外部条件直接影响每一级去噪。"
    )
    add_body(
        doc,
        "式（6）中的“1 + γ”具有明确含义：当 γ=0、β=0 时，FiLM 退化为普通归一化特征，不会在训练初期强行改变主干网络；随着训练进行，网络再学习不同条件下各通道应放大、抑制或平移多少。换言之，FiLM 学到的是条件映射 (t,f,k,p) → (γ,β)，同一个带噪状态 xₜ 在不同日前预测和日历条件下会得到不同的去噪方向。"
    )
    add_heading(doc, "3.4 条件扩散的训练与反向生成", 2)
    add_body(
        doc,
        "训练阶段先向标准化残差 x₀ 逐步加入高斯噪声。由于累积噪声过程具有闭式表达，可在任意时间步 t 直接构造 xₜ："
    )
    add_equation(
        doc,
        "q(xₜ | x₀) = 𝒩(√ᾱₜ x₀, (1 − ᾱₜ)I)，    xₜ = √ᾱₜ x₀ + √(1 − ᾱₜ) ε",
        7,
        "ᾱₜ 为截至 t 的累计保真系数，ε ~ 𝒩(0,I)。t 越大，xₜ 中原始残差信息越少。",
    )
    add_body(
        doc,
        "去噪网络学习在给定 xₜ、扩散步 t 和条件 c 时预测所加入的噪声。三种模型的主要区别，就体现在条件 c 的构造与注入方式。"
    )
    add_equation(
        doc,
        "L_simple(θ) = 𝔼ₓ₀,ₜ,ε [ ‖ε − ε_θ(xₜ, t, c)‖₂² ]",
        8,
        "V4-RS 用输入拼接形成 c；V5-T 仅使用 t；V5-TF 同时使用 t、forecast、calendar 和相对位置。",
    )
    add_body(
        doc,
        "生成阶段从标准高斯噪声 x_T 开始，按 posterior 方差逐步采样 x_{t−1}，直到得到残差场景 x₀。随后通过式（1）还原实际功率场景。"
    )
    add_equation(
        doc,
        "xₜ₋₁ = μ_θ(xₜ,t,c) + σₜ z，    z ~ 𝒩(0,I)",
        9,
        "本算例使用 500 个反向扩散步；随机项 z 使同一条件下能够生成多条具有差异的场景轨迹。",
    )
    add_table(
        doc,
        ["维度", "V4-RS", "V5-T", "V5-TF"],
        [
            ["状态输入", "x_t + 条件拼接，共 14 通道", "仅 x_t，3 通道", "仅 x_t，3 通道"],
            ["扩散时间步 t", "未独立显式注入", "正弦嵌入 + MLP", "正弦嵌入 + MLP"],
            ["日前预测", "输入拼接", "不使用", "独立编码，多尺度 FiLM"],
            ["日历/时间条件", "输入拼接", "不使用", "独立编码，多尺度 FiLM"],
            ["条件作用位置", "U-Net 输入端", "所有残差块", "编码器/瓶颈/解码器"],
            ["参数量", "2.000 M", "1.734 M", "2.532 M"],
            ["实验角色", "旧版基线", "时间步消融", "完整候选模型"],
        ],
        widths=[1.1, 1.7, 1.5, 1.8],
        highlight_rows={6: PALE_GREEN},
    )

    add_heading(doc, "4. 训练收敛与 Checkpoint 选择", 1)
    add_picture(doc, "fig01_training_curves.png", "图 1　三种架构的训练/验证损失与 Rank-1 checkpoint", 6.35)
    add_table(
        doc,
        ["模型", "Rank-1 epoch", "验证 ε-MSE", "训练耗时", "生成耗时"],
        [
            [
                name,
                rank1[name]["checkpoint_epoch"],
                fmt(rank1[name]["validation_epsilon_mse"], 5),
                f"{float(rank1[name]['training_seconds'])/60:.1f} min",
                f"{float(rank1[name]['generation_seconds'])/60:.1f} min",
            ]
            for name in ["V4-RS", "V5-T", "V5-TF"]
        ],
        widths=[1.1, 1.1, 1.4, 1.2, 1.2],
        highlight_rows={2: PALE_GREEN},
    )
    add_body(
        doc,
        "V5-TF 在较早的 epoch 8 即达到最低验证 ε-MSE（0.13055），优于 V5-T 的 0.14264 和 V4-RS 的 0.14627。训练损失继续下降并不等价于场景质量持续提升，因此本阶段坚持使用验证指标选择 checkpoint，而不是直接取最后一轮。"
    )

    add_heading(doc, "5. Rank-1 完整对比结果", 1)
    add_callout(
        doc,
        "主表口径",
        "以下均为物理投影后的验证集结果。CRPS、MVES、ACF 误差、净负荷 MAE 和爬坡 MAE 均为越低越好；90% 覆盖率越接近 90% 越好。",
        fill=LIGHT_BLUE,
    )
    main_rows = []
    for name in ["V4-RS", "V5-T", "V5-TF"]:
        r = rank1[name]
        main_rows.append(
            [
                name,
                fmt(r["constrained_total_crps"]),
                fmt(r["constrained_multivariate_es"], 0),
                fmt(r["constrained_total_acf_mae"], 5),
                fmt(r["constrained_total_coverage_90_pct"]),
                fmt(r["constrained_total_width_90_pct"]),
                fmt(r["constrained_net_load_mae_mw"], 0),
                fmt(r["constrained_net_load_ramp_6h_mae_mw"], 0),
            ]
        )
    add_table(
        doc,
        ["模型", "CRPS", "MVES", "ACF误差", "90%覆盖/%", "90%宽度", "净负荷MAE", "6h爬坡MAE"],
        main_rows,
        widths=[0.72, 0.72, 0.78, 0.72, 0.82, 0.72, 0.86, 0.88],
        highlight_rows={2: PALE_GREEN},
    )
    base, best = rank1["V4-RS"], rank1["V5-TF"]
    improvements = [
        ["CRPS", fmt(base["constrained_total_crps"]), fmt(best["constrained_total_crps"]), f"{pct_reduction(float(base['constrained_total_crps']), float(best['constrained_total_crps'])):.2f}%"],
        ["MVES", fmt(base["constrained_multivariate_es"], 0), fmt(best["constrained_multivariate_es"], 0), f"{pct_reduction(float(base['constrained_multivariate_es']), float(best['constrained_multivariate_es'])):.2f}%"],
        ["ACF 误差", fmt(base["constrained_total_acf_mae"], 5), fmt(best["constrained_total_acf_mae"], 5), f"{pct_reduction(float(base['constrained_total_acf_mae']), float(best['constrained_total_acf_mae'])):.2f}%"],
        ["净负荷 MAE", fmt(base["constrained_net_load_mae_mw"], 0), fmt(best["constrained_net_load_mae_mw"], 0), f"{pct_reduction(float(base['constrained_net_load_mae_mw']), float(best['constrained_net_load_mae_mw'])):.2f}%"],
        ["6h 爬坡 MAE", fmt(base["constrained_net_load_ramp_6h_mae_mw"], 0), fmt(best["constrained_net_load_ramp_6h_mae_mw"], 0), f"{pct_reduction(float(base['constrained_net_load_ramp_6h_mae_mw']), float(best['constrained_net_load_ramp_6h_mae_mw'])):.2f}%"],
    ]
    add_table(doc, ["指标", "V4-RS", "V5-TF", "相对改善"], improvements, widths=[1.5, 1.35, 1.35, 1.35], highlight_rows={0: PALE_GREEN, 1: PALE_GREEN, 2: PALE_GREEN, 3: PALE_GREEN, 4: PALE_GREEN})
    add_picture(doc, "fig02_metric_ratios.png", "图 2　物理投影后 Rank-1 指标相对 V4-RS 的比值", 6.25)
    add_body(
        doc,
        "V5-TF 在五个主要“越低越好”指标上均优于 V4-RS。V5-T 的 90% 覆盖率最高，但其区间更宽、CRPS/MVES/爬坡误差更差，说明较高覆盖率主要来自分布过宽，而非整体概率质量更好。"
    )

    add_heading(doc, "6. 场景曲线与概率包络", 1)
    add_body(
        doc,
        f"代表性窗口采用预先固定的客观规则选择：净负荷日前预测 MAE 最接近 553 个验证窗口的中位数。选中的是第 {fig_summary['representative_window_number_one_based']} 个窗口（零基索引 {fig_summary['representative_window_index_zero_based']}），其日前净负荷 MAE 为 {fig_summary['representative_window_forecast_netload_mae_mw']:.2f} MW。该规则用于展示典型而非极端样本。"
    )
    add_picture(doc, "fig03_envelope_comparison.png", "图 3　三种模型在同一代表性窗口上的 90% 场景包络", 6.45)
    add_picture(doc, "fig04_v5tf_scenario_curves.png", "图 4　V5-TF Rank-1 的 20 条场景曲线、90% 包络和中位数", 6.35)
    add_body(
        doc,
        "三种模型均能跟随风电、光伏和负荷的主趋势。V5-T 在风电和负荷部分时段表现出偏宽包络；V5-TF 的场景中位数整体更贴近实测，同时保留必要离散度。光伏的昼夜周期在投影后严格非负，近夜间点被置零。"
    )

    add_heading(doc, "7. 概率校准与变量分项结果", 1)
    add_picture(doc, "fig06_calibration.png", "图 5　物理投影后的 80%/90%/95% 区间校准曲线", 4.8)
    coverage_rows = []
    for name in ["V4-RS", "V5-T", "V5-TF"]:
        m = cm[name]
        coverage_rows.append(
            [
                name,
                fmt(m["total_coverage_80%"]),
                fmt(m["total_coverage_90%"]),
                fmt(m["total_coverage_95%"]),
                fmt(m["total_width_90%"]),
            ]
        )
    add_table(doc, ["模型", "80%覆盖", "90%覆盖", "95%覆盖", "90%宽度"], coverage_rows, widths=[1.2, 1.2, 1.2, 1.2, 1.2], highlight_rows={2: PALE_GREEN})
    add_body(
        doc,
        "三种模型均存在不同程度的欠覆盖。V5-TF 的 90% 实测覆盖率为 81.60%，比 V4-RS 高 5.12 个百分点，但仍低于名义 90%。因此 V5-TF 已显著改善概率质量，却仍需要后续校准（例如验证集上的分位数扩张或 conformal 校准），不能把当前包络解释为已经完全校准。"
    )
    channel_rows = []
    for name in ["V4-RS", "V5-T", "V5-TF"]:
        m = cm[name]
        channel_rows.append(
            [
                name,
                fmt(m["wind_crps"]),
                fmt(m["solar_crps"]),
                fmt(m["load_crps"]),
                fmt(m["wind_acf_mae"], 5),
                fmt(m["solar_acf_mae"], 5),
                fmt(m["load_acf_mae"], 5),
            ]
        )
    add_table(
        doc,
        ["模型", "风电CRPS", "光伏CRPS", "负荷CRPS", "风电ACF", "光伏ACF", "负荷ACF"],
        channel_rows,
        widths=[0.85, 0.9, 0.9, 0.9, 0.82, 0.82, 0.82],
        highlight_rows={2: PALE_GREEN},
    )
    add_body(
        doc,
        "V5-TF 的总 CRPS 优势主要来自光伏和负荷，同时其风电 CRPS 也优于 V4-RS。分项 ACF 显示 V5-TF 在光伏和负荷上拟合较好；风电的短时相关结构仍有进一步优化空间。"
    )

    add_heading(doc, "8. 时间相关性与爬坡特性", 1)
    add_body(
        doc,
        "时间相关性不是单点误差，而是场景轨迹的连续性和周期性。这里采用三类参数：① ACF MAE，衡量 1–24 h 自相关函数与实测的偏差；② 1 h/6 h 净负荷爬坡 MAE，衡量相邻或跨 6 h 变化量；③ cross-variable correlation MAE，衡量风、光、荷三变量相关矩阵的偏差。"
    )
    add_equation(
        doc,
        "ρ(k) = Σₜ₌₁ᴸ⁻ᵏ (xₜ−x̄)(xₜ₊ₖ−x̄) / Σₜ₌₁ᴸ (xₜ−x̄)²",
        10,
        "ρ(k) 为滞后 k 小时的自相关系数；本报告比较 k=1,…,24 的场景平均 ACF 与实测 ACF。",
    )
    add_equation(
        doc,
        "ACF-MAE = (1/K) Σₖ₌₁ᴷ |ρ_gen(k) − ρ_real(k)|，    K=24",
        11,
        "ACF-MAE 越低，说明生成轨迹的短期记忆和日内周期越接近真实序列。",
    )
    add_equation(
        doc,
        "nₜ = loadₜ − windₜ − solarₜ，    Δₖnₜ = nₜ₊ₖ − nₜ",
        12,
        "nₜ 为净负荷；分别令 k=1 和 k=6，评价 1 h 与 6 h 爬坡变化。",
    )
    add_picture(doc, "fig05_acf_curves.png", "图 6　风电、光伏、负荷的平均自相关函数（0–24 h）", 6.35)
    temporal_rows = []
    for name in ["V4-RS", "V5-T", "V5-TF"]:
        r = rank1[name]
        temporal_rows.append(
            [
                name,
                fmt(r["constrained_total_acf_mae"], 5),
                fmt(r["net_load_ramp_1h_mae_mw"], 0),
                fmt(r["constrained_net_load_ramp_6h_mae_mw"], 0),
                fmt(r["constrained_cross_variable_corr_mae"], 5),
            ]
        )
    add_table(
        doc,
        ["模型", "ACF误差", "净负荷1h爬坡MAE", "净负荷6h爬坡MAE", "跨变量相关误差"],
        temporal_rows,
        widths=[1.05, 1.05, 1.55, 1.55, 1.45],
        highlight_rows={2: PALE_GREEN},
    )
    add_body(
        doc,
        "V5-TF 的 ACF 误差最低，为 0.01096，说明其整体时间依赖最接近实测；6 h 净负荷爬坡误差也低于 V4-RS。需要如实保留的不足是：跨变量相关误差方面 V4-RS（0.04346）仍优于 V5-TF（0.06487），而 V5-T 最差（0.18195）。因此下一阶段应在保留 FiLM 优势的同时加强风—光—荷联合相关结构。"
    )

    add_heading(doc, "9. FiLM 条件消融：新增信息是否真正有效", 1)
    ab_calendar = next(r for r in rows if r["architecture"] == "v5_tf" and r["condition_ablation"] == "calendar")
    ab_forecast = next(r for r in rows if r["architecture"] == "v5_tf" and r["condition_ablation"] == "forecast")
    full = rank1["V5-TF"]
    add_table(
        doc,
        ["V5-TF 设置", "CRPS（原始）", "相对完整模型", "ACF误差", "跨变量相关误差"],
        [
            ["完整 forecast + calendar", fmt(full["total_crps"]), "基准", fmt(full["total_acf_mae"], 5), fmt(full["cross_variable_corr_mae"], 5)],
            ["去掉 calendar", fmt(ab_calendar["total_crps"]), f"+{(float(ab_calendar['total_crps'])/float(full['total_crps'])-1)*100:.2f}%", fmt(ab_calendar["total_acf_mae"], 5), fmt(ab_calendar["cross_variable_corr_mae"], 5)],
            ["去掉 forecast", fmt(ab_forecast["total_crps"]), f"+{(float(ab_forecast['total_crps'])/float(full['total_crps'])-1)*100:.2f}%", fmt(ab_forecast["total_acf_mae"], 5), fmt(ab_forecast["cross_variable_corr_mae"], 5)],
        ],
        widths=[1.7, 1.2, 1.2, 1.2, 1.45],
        highlight_rows={0: PALE_GREEN},
    )
    add_picture(doc, "fig07_ablation_projection.png", "图 7　V5-TF 条件消融与物理投影效果", 6.15)
    add_body(
        doc,
        "去掉 calendar 后 CRPS 上升 11.06%，说明日历/周期信息有稳定贡献；去掉 forecast 后 CRPS 上升 38.54%，ACF 与跨变量相关性也显著恶化，说明日前预测是最关键的外部条件。该消融同时证明 FiLM 不是形式上的结构增加，而是在有效利用条件信息。"
    )

    add_heading(doc, "10. 物理越界来源与最稳妥处理", 1)
    add_heading(doc, "10.1 为什么清洗后仍会生成负太阳能", 2)
    add_body(
        doc,
        "数据清洗约束的是训练样本，不会自动约束生成器的输出空间。模型学习的是标准化残差，而扩散采样从高斯噪声开始，网络输出在数学上仍属于无界实数。夜间 forecast 接近 0，只要生成残差略偏正，通过 actual = forecast − residual 重构后就会得到负光伏。因此负值的根源是“无界随机生成 + 残差重构”，不是原始数据里仍残留负值。"
    )
    add_heading(doc, "10.2 当前采用的物理可行域投影", 2)
    add_callout(
        doc,
        "约束规则",
        "wind = clip(wind, 0, wind_capacity)；solar = clip(solar, 0, solar_capacity)；load = max(load, 0)；当 solar_forecast ≤ 1 MW 时 solar = 0。",
        fill=LIGHT_BLUE,
    )
    add_equation(
        doc,
        "P_w(z) = min(max(z,0), C_w)，    P_l(z) = max(z,0)",
        13,
        "C_w 为风电容量上界；负荷仅施加非负下界。",
    )
    add_equation(
        doc,
        "P_s(z;f_s) = 0，若 f_s ≤ 1 MW；否则 P_s(z;f_s) = min(max(z,0), C_s)",
        14,
        "C_s 为光伏容量上界；f_s≤1 MW 用作当前数据上的近似夜间标志。",
    )
    projection_rows = []
    for name in ["V4-RS", "V5-T", "V5-TF"]:
        r = rank1[name]
        projection_rows.append(
            [
                name,
                fmt(r["any_physical_violation_pct"]),
                fmt(r["constrained_any_physical_violation_pct"]),
                fmt(r["total_crps"]),
                fmt(r["constrained_total_crps"]),
                fmt(r["solar_crps"]),
                fmt(cm[name]["solar_crps"]),
            ]
        )
    add_table(
        doc,
        ["模型", "原始越界/%", "投影后/%", "原始CRPS", "投影后CRPS", "原始光伏CRPS", "投影后光伏CRPS"],
        projection_rows,
        widths=[0.85, 0.9, 0.8, 0.9, 0.95, 1.05, 1.05],
        highlight_rows={2: PALE_GREEN},
    )
    add_body(
        doc,
        "V5-TF 原始任一物理越界率为 19.12%，主要由光伏负值（18.64%）构成。投影后越界率为 0%，总 CRPS 从 732.55 降至 728.39，光伏 CRPS 从 381.22 降至 368.91。原始数组未覆盖，约束后数组与指标单独保存，因此既满足业务可用性，也保留科研审计链。"
    )

    add_heading(doc, "11. 配对统计与 checkpoint 稳定性", 1)
    b_v4 = bootstrap["V5-TF_vs_V4-RS"]
    b_v5t = bootstrap["V5-TF_vs_V5-T"]
    add_table(
        doc,
        ["比较", "窗口平均 CRPS 差值", "95% Bootstrap CI", "V5-TF 更优窗口占比"],
        [
            ["V5-TF − V4-RS", fmt(b_v4["mean_delta"]), f"[{fmt(b_v4['ci95'][0])}, {fmt(b_v4['ci95'][1])}]", f"{b_v4['window_better_pct']:.2f}%"],
            ["V5-TF − V5-T", fmt(b_v5t["mean_delta"]), f"[{fmt(b_v5t['ci95'][0])}, {fmt(b_v5t['ci95'][1])}]", f"{b_v5t['window_better_pct']:.2f}%"],
        ],
        widths=[1.5, 1.55, 1.65, 1.65],
        highlight_rows={0: PALE_GREEN, 1: PALE_GREEN},
    )
    add_body(
        doc,
        "差值小于 0 表示 V5-TF 更好。10,000 次窗口级配对 Bootstrap 的两个置信区间均完全小于 0，说明 V5-TF 的优势不是由少数窗口偶然造成。V5-TF 的 Rank-1（epoch 8）与 Rank-2（epoch 9）总体 CRPS 接近，但 Rank-1 的 ACF 误差和 90% 覆盖率更均衡，因此推荐 epoch 8 作为当前主 checkpoint。"
    )

    add_heading(doc, "12. 结论、限制与下一步", 1)
    conclusions = [
        "可以进入结果分析和项目汇报，无需因为物理投影重新训练。投影是推理/后处理约束，不改变已训练权重。",
        "当前最优路线是 V5-TF Rank-1（epoch 8）。它在 CRPS、MVES、ACF、净负荷和 6 h 爬坡等综合指标上优于 V4-RS。",
        "V5-T 的结果证明：只补扩散时间步、却去掉 forecast/calendar，不足以形成有效升级；V5-TF 的优势来自时间步与外部条件的联合 FiLM 调制。",
        "物理投影解决了负风电/负光伏等越界，且原始结果被完整保留，符合工程应用和科研审计的双重要求。",
        "当前 90% 覆盖率仍偏低，跨变量相关性也未全面超过 V4-RS；这是下一阶段应优先优化的两项。",
        "正式对外定版前，建议增加至少 3 个训练随机种子，按预先固定规则选模型，再一次性评估测试集，避免反复查看测试集造成信息泄漏。",
    ]
    for i, text in enumerate(conclusions, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.16)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{i}. {text}")
        set_run_font(r, 11.5)

    add_callout(
        doc,
        "推荐的汇报表述",
        "“V5-TF 在验证集 Stage-1 统一算例中取得当前最佳综合结果，物理投影后零越界；相比 V4-RS，概率评分、时间相关性和净负荷指标均明显改善。下一阶段将通过多随机种子和最终测试集验证确认泛化稳定性。”",
    )

    add_heading(doc, "附录 A　结果文件与复现口径", 1)
    add_table(
        doc,
        ["内容", "文件/目录"],
        [
            ["主对比表", r"outputs_shandong\v5_stage1\comparisons\20260723_physical_projection\v5_stage1_comparison.csv"],
            ["原始场景", "各 result_dir 下 actual_scenarios.npy（保持不变）"],
            ["约束场景", "各 result_dir 下 actual_scenarios_constrained.npy"],
            ["约束指标", "各 result_dir 下 metrics_constrained.json"],
            ["投影审计", "各 result_dir 下 physical_projection.json"],
            ["V5-TF 推荐 checkpoint", "training run 20260722_155013... / model_epoch_8.pt"],
        ],
        widths=[1.45, 4.7],
    )
    add_body(
        doc,
        "说明：本报告所有主结论来自 val split；场景数、采样种子、采样方差类型、反向步数及代码提交均已固定。报告中的曲线由保存的 NPY 数组重新计算和绘制，表格由 CSV/JSON 自动读取，避免手工抄录。"
    )

    add_heading(doc, "附录 B　主要概率指标的数学含义", 1)
    add_body(
        doc,
        "设 S 为场景数，x⁽ˢ⁾ 为第 s 条生成场景，y 为实测值。CRPS 同时考虑场景与实测的偏差以及场景集合自身的离散程度："
    )
    add_equation(
        doc,
        "CRPS = (1/S)Σₛ|x⁽ˢ⁾−y| − (1/2S²)ΣₛΣₛ′|x⁽ˢ⁾−x⁽ˢ′⁾|",
        15,
        "第一项惩罚不准确，第二项避免所有场景坍缩到同一点；综合值越低越好。",
    )
    add_body(
        doc,
        "Energy Score 是 CRPS 的多变量推广。将一个 168 h 风—光—荷联合轨迹视为向量后，可写为："
    )
    add_equation(
        doc,
        "ES = (1/S)Σₛ‖x⁽ˢ⁾−y‖₂ − (1/2S²)ΣₛΣₛ′‖x⁽ˢ⁾−x⁽ˢ′⁾‖₂",
        16,
        "本报告中的 MVES 用于评价多变量联合场景，越低表示联合分布整体更接近实测。",
    )
    add_body(
        doc,
        "覆盖率用于检查概率包络是否校准。名义覆盖率为 α 时，定义为："
    )
    add_equation(
        doc,
        "Coverage_α = (1/N)Σᵢ 𝟙{q_(1−α)/2,i ≤ yᵢ ≤ q_(1+α)/2,i}",
        17,
        "Coverage_α 应接近 α；但还必须同时观察区间宽度，不能只靠无限扩大包络提高覆盖率。",
    )

    # Document properties and final save
    doc.core_properties.title = "V5条件扩散模型Stage-1对比算例与消融实验报告"
    doc.core_properties.subject = "V4-RS、V5-T、V5-TF 验证集对比与物理可行性审计"
    doc.core_properties.author = "DM 项目组"
    doc.core_properties.keywords = "条件扩散, FiLM, 场景生成, CRPS, ACF, 物理投影"
    doc.save(OUTPUT)
    print(OUTPUT)
    print("size", OUTPUT.stat().st_size)
    print("bootstrap", json.dumps(bootstrap, ensure_ascii=False))


if __name__ == "__main__":
    main()
