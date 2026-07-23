from pathlib import Path
import json
import zipfile

from docx import Document
from docx.oxml.ns import qn


path = Path(r"D:\DM_local\reports\v5_stage1_report\_work\reference.docx")
doc = Document(path)


def run_info(run):
    rpr = run._element.rPr
    rfonts = rpr.rFonts if rpr is not None else None
    return {
        "text": run.text,
        "font": run.font.name,
        "east_asia": rfonts.get(qn("w:eastAsia")) if rfonts is not None else None,
        "size_pt": run.font.size.pt if run.font.size else None,
        "bold": run.bold,
        "italic": run.italic,
        "color": str(run.font.color.rgb) if run.font.color.rgb else None,
    }


paragraphs = []
for index, paragraph in enumerate(doc.paragraphs):
    fmt = paragraph.paragraph_format
    paragraphs.append({
        "index": index,
        "style": paragraph.style.name,
        "text": paragraph.text,
        "alignment": str(paragraph.alignment),
        "space_before_pt": fmt.space_before.pt if fmt.space_before else None,
        "space_after_pt": fmt.space_after.pt if fmt.space_after else None,
        "line_spacing": fmt.line_spacing,
        "left_indent_in": fmt.left_indent.inches if fmt.left_indent else None,
        "first_line_indent_in": (
            fmt.first_line_indent.inches if fmt.first_line_indent else None
        ),
        "keep_with_next": fmt.keep_with_next,
        "page_break_before": fmt.page_break_before,
        "runs": [run_info(run) for run in paragraph.runs],
    })

tables = []
for table_index, table in enumerate(doc.tables):
    tables.append({
        "index": table_index,
        "style": table.style.name if table.style else None,
        "rows": len(table.rows),
        "cols": len(table.columns),
        "cells": [[cell.text for cell in row.cells] for row in table.rows],
    })

styles = {}
for name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "List Paragraph"):
    try:
        style = doc.styles[name]
    except KeyError:
        continue
    fmt = style.paragraph_format
    styles[name] = {
        "font": style.font.name,
        "size_pt": style.font.size.pt if style.font.size else None,
        "bold": style.font.bold,
        "color": str(style.font.color.rgb) if style.font.color.rgb else None,
        "alignment": str(fmt.alignment),
        "space_before_pt": fmt.space_before.pt if fmt.space_before else None,
        "space_after_pt": fmt.space_after.pt if fmt.space_after else None,
        "line_spacing": fmt.line_spacing,
        "first_line_indent_in": (
            fmt.first_line_indent.inches if fmt.first_line_indent else None
        ),
    }

with zipfile.ZipFile(path) as archive:
    package = [
        {
            "path": info.filename,
            "size": info.file_size,
            "compressed": info.compress_size,
        }
        for info in archive.infolist()
    ]

output = {
    "paragraphs": paragraphs,
    "tables": tables,
    "styles": styles,
    "package": package,
}
Path(r"D:\DM_local\reports\v5_stage1_report\_work\reference_evidence.json").write_text(
    json.dumps(output, ensure_ascii=False, indent=2),
    encoding="utf-8",
)
print("paragraphs", len(paragraphs), "tables", len(tables), "package_parts", len(package))
for item in paragraphs:
    print(item["index"], item["style"], repr(item["text"][:120]))
for item in tables:
    print("TABLE", item["index"], item["style"], item["rows"], item["cols"])
