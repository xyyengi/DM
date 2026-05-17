#!/usr/bin/env python3
"""
扩散模型 vs CGAN 指标对比分析
生成可视化图表和周报数据
"""

import json
import numpy as np
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # 非交互式后端

# 设置中文字体
plt.rcParams['font.sans-serif'] = ['DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

# 扩散模型指标 (guidance=0, 最佳配置)
diffusion_metrics = {
    "multivariate_es": 2.4309841103023957,
    "wind_crps": 0.14187952217196736,
    "wind_energy_score": -4.37550040880839,
    "wind_coverage_100%": 67.2875682516084,
    "wind_width_100%": 31.364739370543145,
    "wind_coverage_90%": 59.91108485765613,
    "wind_width_90%": 26.607314955064908,
    "wind_coverage_80%": 50.70718400722117,
    "wind_width_80%": 21.8498905396151,
    "wind_acf_mae": 0.2053708281073341,
    "solar_crps": 0.044903617761142704,
    "solar_energy_score": -1.8069213973151312,
    "solar_coverage_100%": 89.10983088645233,
    "solar_width_100%": 20.384278482091727,
    "solar_coverage_90%": 85.15444553588021,
    "solar_width_90%": 16.80575511847124,
    "solar_coverage_80%": 77.31756798612402,
    "solar_width_80%": 13.227231755272564,
    "solar_acf_mae": 0.07462907943408936,
    "load_crps": 0.024197086126540705,
    "load_energy_score": -0.8078861713409426,
    "load_coverage_100%": 91.72739621773258,
    "load_width_100%": 43.50836299362246,
    "load_coverage_90%": 88.45861983522269,
    "load_width_90%": 35.13514767638116,
    "load_coverage_80%": 81.21521048486296,
    "load_width_80%": 26.761932359557978,
    "load_acf_mae": 0.12949474616087475,
    "total_crps": 0.07032674201988359,
    "total_energy_score": -2.3301026591548215,
    "total_coverage_100%": 82.70826511859777,
    "total_width_100%": 31.752460282085778,
    "total_acf_mae": 0.13649821790076608
}

# CGAN指标 (label1)
cgan_metrics = {
    "multivariate_es": 2.3541811651410516,
    "wind_crps": 0.09945054641972247,
    "wind_energy_score": -3.1725577958690314,
    "wind_coverage_100%": 77.22663139329806,
    "wind_width_100%": 51.13587041565493,
    "wind_coverage_90%": 70.2932098765432,
    "wind_width_90%": 44.03417361707513,
    "wind_coverage_80%": 61.155202821869494,
    "wind_width_80%": 36.93247681849535,
    "wind_acf_mae": 0.016359370323516084,
    "solar_crps": 0.06454539913972981,
    "solar_energy_score": -2.621566206067952,
    "solar_coverage_100%": 77.49118165784833,
    "solar_width_100%": 28.556711245618082,
    "solar_coverage_90%": 69.34523809523809,
    "solar_width_90%": 25.36429017123017,
    "solar_coverage_80%": 61.00088183421517,
    "solar_width_80%": 22.17186909684227,
    "solar_acf_mae": 0.03124714912976398,
    "load_crps": 0.06962767752507104,
    "load_energy_score": -2.053778598972853,
    "load_coverage_100%": 79.65167548500882,
    "load_width_100%": 33.30509995508077,
    "load_coverage_90%": 72.13403880070547,
    "load_width_90%": 28.755748307965284,
    "load_coverage_80%": 63.75661375661375,
    "load_width_80%": 24.206396660849816,
    "load_acf_mae": 0.010387479072721419,
    "total_crps": 0.07787454102817444,
    "total_energy_score": -2.615967533636612,
    "total_coverage_100%": 78.12316284538507,
    "total_width_100%": 37.66589387211793,
    "total_acf_mae": 0.019331332842000493
}

def calculate_improvement(diff_val, cgan_val, higher_is_better=True):
    """计算改进百分比"""
    if higher_is_better:
        return ((diff_val - cgan_val) / cgan_val) * 100
    else:
        return ((cgan_val - diff_val) / cgan_val) * 100

def create_comparison_plots():
    """创建对比可视化图表"""
    
    # 创建大图
    fig = plt.figure(figsize=(20, 12))
    
    # 1. Coverage对比 (100%, 90%, 80%)
    ax1 = plt.subplot(2, 3, 1)
    categories = ['Wind', 'Solar', 'Load', 'Total']
    coverage_types = ['100%', '90%', '80%']
    
    x = np.arange(len(categories))
    width = 0.25
    
    # 颜色与样式配置
    diff_color = '#1f77b4'
    cgan_color = '#ff7f0e'
    hatches = ['','//','\\']

    for i, cov_type in enumerate(coverage_types):
        diff_vals = [diffusion_metrics.get(f'{cat.lower()}_coverage_{cov_type}', np.nan)
                     for cat in ['wind', 'solar', 'load', 'total']]
        cgan_vals = [cgan_metrics.get(f'{cat.lower()}_coverage_{cov_type}', np.nan)
                     for cat in ['wind', 'solar', 'load', 'total']]

        b1 = ax1.bar(x - width/2 + i*width, diff_vals, width,
                     label=f'Diffusion {cov_type}', color=diff_color,
                     edgecolor='k', linewidth=0.4, alpha=0.9, hatch=hatches[i])
        b2 = ax1.bar(x + width/2 + i*width, cgan_vals, width,
                     label=f'CGAN {cov_type}', color=cgan_color,
                     edgecolor='k', linewidth=0.4, alpha=0.9, hatch=hatches[i])

    # 去重图例（保留出现顺序）
    handles, labels = ax1.get_legend_handles_labels()
    from collections import OrderedDict
    by_label = OrderedDict()
    for h, l in zip(handles, labels):
        if l not in by_label:
            by_label[l] = h
    ax1.legend(by_label.values(), by_label.keys(), fontsize='small')
    
    ax1.set_ylabel('Coverage (%)')
    ax1.set_title('Coverage Comparison Across Models')
    ax1.set_xticks(x)
    ax1.set_xticklabels(categories)
    ax1.legend()
    ax1.grid(True, alpha=0.3)
    
    # 2. Width对比
    ax2 = plt.subplot(2, 3, 2)
    width_types = ['100%', '90%', '80%']

    for i, w_type in enumerate(width_types):
        diff_vals = [diffusion_metrics.get(f'{cat.lower()}_width_{w_type}', np.nan)
                     for cat in ['wind', 'solar', 'load', 'total']]
        cgan_vals = [cgan_metrics.get(f'{cat.lower()}_width_{w_type}', np.nan)
                     for cat in ['wind', 'solar', 'load', 'total']]

        ax2.bar(x - width/2 + i*width, diff_vals, width,
                label=f'Diffusion {w_type}', color=diff_color,
                edgecolor='k', linewidth=0.4, alpha=0.9, hatch=hatches[i])
        ax2.bar(x + width/2 + i*width, cgan_vals, width,
                label=f'CGAN {w_type}', color=cgan_color,
                edgecolor='k', linewidth=0.4, alpha=0.9, hatch=hatches[i])

    # 去重图例（保留出现顺序）
    handles, labels = ax2.get_legend_handles_labels()
    from collections import OrderedDict
    by_label = OrderedDict()
    for h, l in zip(handles, labels):
        if l not in by_label:
            by_label[l] = h
    ax2.legend(by_label.values(), by_label.keys(), fontsize='small')

    ax2.set_ylabel('Width')
    ax2.set_title('Width Comparison Across Models')
    ax2.set_xticks(x)
    ax2.set_xticklabels(categories)
    ax2.grid(True, alpha=0.3)
    
    # 3. CRPS对比
    ax3 = plt.subplot(2, 3, 3)
    crps_categories = ['Wind', 'Solar', 'Load', 'Total']
    diff_crps = [diffusion_metrics[f'{cat.lower()}_crps'] 
                 for cat in ['wind', 'solar', 'load', 'total']]
    cgan_crps = [cgan_metrics[f'{cat.lower()}_crps'] 
                 for cat in ['wind', 'solar', 'load', 'total']]
    
    x_crps = np.arange(len(crps_categories))
    ax3.bar(x_crps - width/2, diff_crps, width, label='Diffusion', alpha=0.8)
    ax3.bar(x_crps + width/2, cgan_crps, width, label='CGAN', alpha=0.8)
    
    ax3.set_ylabel('CRPS (lower is better)')
    ax3.set_title('CRPS Comparison')
    ax3.set_xticks(x_crps)
    ax3.set_xticklabels(crps_categories)
    ax3.legend()
    ax3.grid(True, alpha=0.3)
    
    # 4. Energy Score对比
    ax4 = plt.subplot(2, 3, 4)
    diff_es = [diffusion_metrics[f'{cat.lower()}_energy_score'] 
               for cat in ['wind', 'solar', 'load', 'total']]
    cgan_es = [cgan_metrics[f'{cat.lower()}_energy_score'] 
               for cat in ['wind', 'solar', 'load', 'total']]
    
    ax4.bar(x_crps - width/2, diff_es, width, label='Diffusion', alpha=0.8)
    ax4.bar(x_crps + width/2, cgan_es, width, label='CGAN', alpha=0.8)
    
    ax4.set_ylabel('Energy Score (higher is better)')
    ax4.set_title('Energy Score Comparison')
    ax4.set_xticks(x_crps)
    ax4.set_xticklabels(crps_categories)
    ax4.legend()
    ax4.grid(True, alpha=0.3)
    
    # 5. ACF MAE对比
    ax5 = plt.subplot(2, 3, 5)
    diff_acf = [diffusion_metrics[f'{cat.lower()}_acf_mae'] 
                for cat in ['wind', 'solar', 'load', 'total']]
    cgan_acf = [cgan_metrics[f'{cat.lower()}_acf_mae'] 
                for cat in ['wind', 'solar', 'load', 'total']]
    
    ax5.bar(x_crps - width/2, diff_acf, width, label='Diffusion', alpha=0.8)
    ax5.bar(x_crps + width/2, cgan_acf, width, label='CGAN', alpha=0.8)
    
    ax5.set_ylabel('ACF MAE (lower is better)')
    ax5.set_title('ACF MAE Comparison')
    ax5.set_xticks(x_crps)
    ax5.set_xticklabels(crps_categories)
    ax5.legend()
    ax5.grid(True, alpha=0.3)
    
    # 6. 综合雷达图
    ax6 = plt.subplot(2, 3, 6, projection='polar')
    
    # 选择关键指标进行雷达图展示
    radar_metrics = ['total_coverage_100%', 'total_crps', 'total_energy_score', 'total_acf_mae']
    radar_labels = ['Coverage', 'CRPS', 'Energy Score', 'ACF MAE']
    
    # 归一化到0-1范围
    diff_radar = []
    cgan_radar = []
    
    for metric in radar_metrics:
        diff_val = diffusion_metrics[metric]
        cgan_val = cgan_metrics[metric]
        
        # 对于CRPS和ACF MAE，越低越好，需要反转
        if 'crps' in metric or 'acf' in metric:
            max_val = max(diff_val, cgan_val) * 1.2
            diff_norm = 1 - (diff_val / max_val)
            cgan_norm = 1 - (cgan_val / max_val)
        else:
            max_val = max(diff_val, cgan_val) * 1.2
            diff_norm = diff_val / max_val
            cgan_norm = cgan_val / max_val
        
        diff_radar.append(diff_norm)
        cgan_radar.append(cgan_norm)
    
    angles = np.linspace(0, 2 * np.pi, len(radar_labels), endpoint=False).tolist()
    diff_radar += diff_radar[:1]
    cgan_radar += cgan_radar[:1]
    angles += angles[:1]
    
    ax6.plot(angles, diff_radar, 'o-', linewidth=2, label='Diffusion', color='blue')
    ax6.fill(angles, diff_radar, alpha=0.25, color='blue')
    ax6.plot(angles, cgan_radar, 'o-', linewidth=2, label='CGAN', color='red')
    ax6.fill(angles, cgan_radar, alpha=0.25, color='red')
    
    ax6.set_xticks(angles[:-1])
    ax6.set_xticklabels(radar_labels)
    ax6.set_ylim(0, 1)
    ax6.set_title('Overall Performance Radar')
    ax6.legend(loc='upper right', bbox_to_anchor=(1.3, 1.0))
    
    # 注释每个子图的条形图数值
    def annotate_ax(ax, fmt="{:.2f}"):
        for p in ax.patches:
            try:
                h = p.get_height()
            except Exception:
                continue
            if np.isnan(h):
                continue
            ax.annotate(fmt.format(h),
                        (p.get_x() + p.get_width() / 2, h),
                        ha='center', va='bottom', fontsize=8, rotation=0)

    for a in [ax1, ax2, ax3, ax4, ax5]:
        annotate_ax(a)

    plt.tight_layout()
    plt.savefig('model_comparison_charts.png', dpi=300, bbox_inches='tight')
    print("✓ 对比图表已保存: model_comparison_charts.png")

    return fig

def generate_report():
    """生成对比分析报告"""
    
    report = []
    report.append("# 扩散模型 vs CGAN 性能对比周报")
    report.append(f"\n生成时间: 2026-05-17")
    report.append(f"对比配置: 扩散模型(guidance=0) vs CGAN(label1)")
    # 将生成的图嵌入报告（Markdown）
    report.append("\n![Model Comparison](model_comparison_charts.png)")
    report.append("\n" + "="*80)
    
    # 关键指标对比
    report.append("\n## 一、关键指标对比")
    report.append("\n### 1. 综合指标")
    report.append(f"\n| 指标 | 扩散模型 | CGAN | 改进率 |")
    report.append(f"|------|----------|------|--------|")
    
    total_metrics = [
        ('total_crps', 'CRPS', False),
        ('total_energy_score', 'Energy Score', True),
        ('total_coverage_100%', 'Coverage 100%', True),
        ('total_width_100%', 'Width 100%', False),
        ('total_acf_mae', 'ACF MAE', False)
    ]
    
    for metric, name, higher_better in total_metrics:
        diff_val = diffusion_metrics[metric]
        cgan_val = cgan_metrics[metric]
        improvement = calculate_improvement(diff_val, cgan_val, higher_better)
        
        arrow = "↑" if improvement > 0 else "↓"
        report.append(f"| {name} | {diff_val:.4f} | {cgan_val:.4f} | {improvement:+.1f}% {arrow} |")
    
    # 各维度详细对比
    report.append("\n### 2. 各维度详细对比")
    
    for category in ['Wind', 'Solar', 'Load']:
        cat_lower = category.lower()
        report.append(f"\n#### {category}")
        report.append(f"\n| 指标 | 扩散模型 | CGAN | 改进率 |")
        report.append(f"|------|----------|------|--------|")
        
        cat_metrics = [
            (f'{cat_lower}_crps', 'CRPS', False),
            (f'{cat_lower}_energy_score', 'Energy Score', True),
            (f'{cat_lower}_coverage_100%', 'Coverage 100%', True),
            (f'{cat_lower}_width_100%', 'Width 100%', False),
            (f'{cat_lower}_acf_mae', 'ACF MAE', False)
        ]
        
        for metric, name, higher_better in cat_metrics:
            diff_val = diffusion_metrics[metric]
            cgan_val = cgan_metrics[metric]
            improvement = calculate_improvement(diff_val, cgan_val, higher_better)
            
            arrow = "↑" if improvement > 0 else "↓"
            report.append(f"| {name} | {diff_val:.4f} | {cgan_val:.4f} | {improvement:+.1f}% {arrow} |")
    
    # 关键发现
    report.append("\n## 二、关键发现")
    report.append("\n### 优势分析")
    
    # 扩散模型优势
    diff_advantages = []
    if diffusion_metrics['solar_coverage_100%'] > cgan_metrics['solar_coverage_100%']:
        diff_advantages.append(f"光伏Coverage显著优于CGAN ({diffusion_metrics['solar_coverage_100%']:.1f}% vs {cgan_metrics['solar_coverage_100%']:.1f}%)")
    if diffusion_metrics['load_coverage_100%'] > cgan_metrics['load_coverage_100%']:
        diff_advantages.append(f"负荷Coverage显著优于CGAN ({diffusion_metrics['load_coverage_100%']:.1f}% vs {cgan_metrics['load_coverage_100%']:.1f}%)")
    if diffusion_metrics['total_crps'] < cgan_metrics['total_crps']:
        diff_advantages.append(f"总体CRPS更低 ({diffusion_metrics['total_crps']:.4f} vs {cgan_metrics['total_crps']:.4f})")
    
    for adv in diff_advantages:
        report.append(f"- ✓ {adv}")
    
    report.append("\n### 劣势分析")
    
    # 扩散模型劣势
    diff_disadvantages = []
    if diffusion_metrics['wind_coverage_100%'] < cgan_metrics['wind_coverage_100%']:
        diff_disadvantages.append(f"风电Coverage低于CGAN ({diffusion_metrics['wind_coverage_100%']:.1f}% vs {cgan_metrics['wind_coverage_100%']:.1f}%)")
    if diffusion_metrics['wind_crps'] > cgan_metrics['wind_crps']:
        diff_disadvantages.append(f"风电CRPS较高 ({diffusion_metrics['wind_crps']:.4f} vs {cgan_metrics['wind_crps']:.4f})")
    if diffusion_metrics['total_acf_mae'] > cgan_metrics['total_acf_mae']:
        diff_disadvantages.append(f"时间相关性(ACF MAE)较差 ({diffusion_metrics['total_acf_mae']:.4f} vs {cgan_metrics['total_acf_mae']:.4f})")
    
    for dis in diff_disadvantages:
        report.append(f"- ⚠ {dis}")
    
    # 结论
    report.append("\n## 三、结论与建议")
    report.append("\n### 总体评价")
    
    # 计算综合得分
    diff_score = 0
    cgan_score = 0
    
    # Coverage权重
    for cat in ['wind', 'solar', 'load']:
        if diffusion_metrics[f'{cat}_coverage_100%'] > cgan_metrics[f'{cat}_coverage_100%']:
            diff_score += 1
        else:
            cgan_score += 1
    
    # CRPS权重
    if diffusion_metrics['total_crps'] < cgan_metrics['total_crps']:
        diff_score += 2
    else:
        cgan_score += 2
    
    report.append(f"\n- 扩散模型综合得分: {diff_score}")
    report.append(f"- CGAN综合得分: {cgan_score}")
    
    if diff_score > cgan_score:
        report.append(f"\n**结论**: 扩散模型整体表现更优，特别是在光伏和负荷的Coverage方面。")
    else:
        report.append(f"\n**结论**: CGAN整体表现更优，但扩散模型在特定场景（光伏、负荷）有优势。")
    
    report.append("\n### 后续优化建议")
    report.append("1. **风电优化**: 扩散模型在风电Coverage上明显落后，需要针对性优化")
    report.append("2. **时间相关性**: ACF MAE较高，考虑引入时间序列特定的损失函数")
    report.append("3. **宽度平衡**: 扩散模型Width普遍较窄，可能需要调整生成多样性")
    report.append("4. **CGAN对比**: CGAN在风电和整体ACF上表现更好，可借鉴其架构设计")

    # 对消融实验（无条件/弱条件）优于带条件引导的可能解释
    report.append("\n## 四、关于消融实验优于有条件引导的可能原因")
    report.append("- 指导强度偏差：过强的条件引导会把生成样本收敛到训练分布的常见模式，降低样本多样性，导致部分极端或长尾场景（如风电峰谷）被忽视。")
    report.append("- 信号不匹配：条件信息（如气象或外生变量）与模型接受方式不一致，会把模型推向错误的局部最优，反而降低部分指标。")
    report.append("- 数据/任务差异：当前周尺度场景中，风电的不确定性、非平稳性更强，弱化引导或去掉条件可能反而保留了更丰富的生成多样性，从而在总体CRPS和部分覆盖率上表现更好。")
    report.append("\n建议：做一组控制实验，在不同的 guidance 强度（如 0, 0.1, 0.5, 1.0）下评估各项指标，记录不同子群（风/光/负荷）表现，定位是否为指导强度或条件不匹配导致的退化。")

    # 后续计划（用户提供，整理为执行项）
    report.append("\n## 五、后续计划（执行项）")
    report.append("以下工作将并行推进，优先在新分支上开展实验以便回溯与对比：")
    report.append("\n1) 梯度引导符号/强度试验（短期）")
    report.append("- 在新分支上尝试反转或弱化梯度引导符号，测试能否提升风电Coverage（当前约 50.71%）。")
    report.append("- 系统化 sweep：对比正向/负向/0（无引导）及若干强度（小步长），记录风/光/负荷的覆盖率、CRPS、ACF MAE。")

    report.append("2) 流匹配与物理约束（面向山东电网 96 节点扩展，中期）")
    report.append("- 空间拓扑升级：将训练数据从单站点扩展至山东电网 96 节点，引入电网拓扑矩阵，升级主干为时空图神经网络以捕捉多节点间的空间相关性。")
    report.append("- 算法加速（流匹配）：尝试引入 Flow Matching 重构生成轨迹，将去噪步数缩短至 5-10 步以提升高维生成效率。")
    report.append("- 物理约束注入：在去噪过程中注入有功潮流等电力系统方程作为梯度引导项，约束并修正不符合物理规律的生成场景，保证 96 节点的安全边际。")

    report.append("3) 极值与条件控制（针对极端天气）")
    report.append("- 显式条件控制：在条件输入中加入极端天气事件标签（如寒潮、大风切出等），训练模型在这些条件下向更符合长尾分布的样本空间去噪。")
    report.append("- 针对性评估：对极端事件子集做单独评估指标（覆盖率/CRPS/能量分布），确保改动提升尾部性能而不损害总体表现。")

    report.append("\n每项任务都应包含：实验分支、数据准备说明、评估指标清单、期望改进阈值与回退条件。")
    
    return "\n".join(report)

if __name__ == "__main__":
    print("正在生成对比分析...")
    
    # 生成可视化图表
    create_comparison_plots()
    
    # 生成报告
    report = generate_report()
    
    # 保存报告
    with open('weekly_report_diffusion_vs_cgan.md', 'w', encoding='utf-8') as f:
        f.write(report)
    
    print("✓ 周报已保存: weekly_report_diffusion_vs_cgan.md")

    # 尝试生成 Word 文档（需要 python-docx）
    try:
        from docx import Document
        from docx.shared import Inches

        doc = Document()
        doc.add_heading('扩散模型 vs CGAN 性能对比周报', level=1)
        doc.add_paragraph(f'生成时间: 2026-05-17')
        # 插入对比图
        try:
            doc.add_picture('model_comparison_charts.png', width=Inches(6))
        except Exception:
            doc.add_paragraph('（未能插入图片：model_comparison_charts.png）')

        # 将 Markdown 文本简单写入 Word
        with open('weekly_report_diffusion_vs_cgan.md', 'r', encoding='utf-8') as md_f:
            for line in md_f:
                doc.add_paragraph(line.rstrip())

        doc.save('weekly_report_diffusion_vs_cgan.docx')
        print('✓ Word已保存: weekly_report_diffusion_vs_cgan.docx')
    except ImportError:
        print("python-docx 未安装。请运行 `pip install python-docx` 然后重新运行脚本以生成 Word 文档.")
    print("\n" + "="*80)
    print(report)
